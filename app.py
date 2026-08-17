# ============================================================
# app.py - AI Clinical Assistant (Complete MVP) 
# ============================================================

import streamlit as st
import time
import re
import json
import io
import html
from datetime import datetime
from PIL import Image

# ---------- 0. Safely import core module ----------
try:
    from core import run_pipeline
    core_available = True
except ImportError as e:
    core_available = False
    core_error = str(e)

# ---------- 1. Page Config & Optimized Light Theme CSS ----------
st.set_page_config(page_title="AI Clinical Assistant", layout="wide")

if not core_available:
    st.error(f"❌ Core module 'core.py' not found or has missing dependencies: {core_error}")
    st.error("Please ensure 'core.py' is in the same directory and all dependencies are installed.")
    st.stop()

st.markdown("""
<style>
    html, body, .stApp {
        background: linear-gradient(135deg, #f5faff 0%, #ffe6f0 100%) !important;
        background-color: #f5faff !important;
    }
    h1, h2, h3, h4, h5, h6, p, label, .stMarkdown, .stCaption, .stText, .stException {
        color: #2c3e50 !important;
    }
    [data-testid="stCheckbox"] label, [data-testid="stRadio"] label {
        color: #2c3e50 !important;
    }
    [data-testid="stSidebar"] {
        background: rgba(255, 255, 255, 0.9) !important;
        backdrop-filter: blur(8px);
    }
    .stTextArea textarea, .stTextInput input, .stNumberInput input, div[data-baseweb="select"] > div {
        background-color: #ffffff !important;
        color: #2c3e50 !important;
        border-radius: 12px !important;
        border: 1px solid #d0e8f0 !important;
    }
    div[data-baseweb="menu"] * {
        color: #2c3e50 !important;
        background-color: #ffffff !important;
    }
    div[data-baseweb="select"] svg {
        fill: #2c3e50 !important;
    }

    div[data-testid="stButton"] > button[kind="primary"],
    button[data-testid="baseButton-primary"] {
        background-color: #7ec8e3 !important;
        color: #ffffff !important;
        border-radius: 20px !important;
        padding: 0.5rem 1.6rem !important;
        font-weight: bold !important;
        border: none !important;
        box-shadow: 0 4px 10px rgba(126, 200, 227, 0.3) !important;
        transition: all 0.2s !important;
    }
    div[data-testid="stButton"] > button[kind="primary"]:hover {
        background-color: #5ab3d4 !important;
        transform: translateY(-1px) !important;
    }
    div[data-testid="stButton"] > button[kind="primary"] * {
        color: #ffffff !important;
    }

    div[data-testid="stButton"] > button[kind="secondary"],
    button[data-testid="baseButton-secondary"] {
        background-color: #ffffff !important;
        color: #2c3e50 !important;
        border: 1px solid #b0d8e8 !important;
        border-radius: 8px !important;
        padding: 0.25rem 0.5rem !important;
        font-size: 0.88rem !important;
        font-weight: 600 !important;
        box-shadow: 0 1px 3px rgba(0,0,0,0.03) !important;
        transition: all 0.15s !important;
    }
    div[data-testid="stButton"] > button[kind="secondary"]:hover {
        background-color: #e8f4f8 !important;
        border-color: #7ec8e3 !important;
        color: #1b4965 !important;
        transform: scale(1.03) !important;
    }
    div[data-testid="stButton"] > button[kind="secondary"] * {
        color: #2c3e50 !important;
    }

    /* 表单提交按钮（与 Run Pipeline 保持一致） */
    div[data-testid="stFormSubmitButton"] > button {
    	background-color: #7ec8e3 !important;
   	color: #ffffff !important;
    	border-radius: 20px !important;
    	padding: 0.5rem 1.6rem !important;
    	font-weight: bold !important;
    	border: none !important;
    	box-shadow: 0 4px 10px rgba(126, 200, 227, 0.3) !important;
    	transition: all 0.2s !important;
    }
    div[data-testid="stFormSubmitButton"] > button:hover {
    	background-color: #5ab3d4 !important;
    	transform: translateY(-1px) !important;
    }
    div[data-testid="stFormSubmitButton"] > button * {
    	color: #ffffff !important;
    }

    .card {
        background: rgba(255, 255, 255, 0.85) !important;
        backdrop-filter: blur(4px);
        border-radius: 16px;
        padding: 1.2rem;
        margin-bottom: 0.8rem;
        border: 1px solid rgba(255, 255, 255, 0.9);
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    }
    .card, .card p, .card li, .card strong, .card span, .card h1, .card h2, .card h3, .card h4, .card div {
        color: #1e2a3a !important;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: rgba(255,255,255,0.4);
        border-radius: 16px;
        padding: 4px;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 12px;
        padding: 0.5rem 1.5rem;
        background-color: transparent;
        color: #2c3e50 !important;
    }
    .stTabs [aria-selected="true"] {
        background-color: #ffffffcc !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }
    .recovery-banner {
        background-color: #ffecb3;
        border-left: 6px solid #ffb300;
        padding: 12px 20px;
        border-radius: 12px;
        margin-bottom: 16px;
    }
    .stAlert, .stSuccess, .stWarning, .stError, .stInfo {
        border-radius: 12px !important;
    }
</style>
""", unsafe_allow_html=True)

# ---------- 2. Session State Initialization ----------
st.session_state.setdefault("processed", False)
st.session_state.setdefault("file_version", 0)
st.session_state.setdefault("input_text", "")
st.session_state.setdefault("tab1_text_area", "")
st.session_state.setdefault("patient_name", "")
st.session_state.setdefault("age", "Not specified")
st.session_state.setdefault("gender", "Not specified")
st.session_state.setdefault("race", "Not specified")
st.session_state.setdefault("prev_notes", "")

# Draft storage
st.session_state.setdefault("draft_input", "")
st.session_state.setdefault("draft_cleaned", "")
st.session_state.setdefault("draft_diseases", [])
st.session_state.setdefault("draft_soap_dict", {})
st.session_state.setdefault("draft_post", {})
st.session_state.setdefault("draft_patient_name", "")
st.session_state.setdefault("draft_age", "Not specified")
st.session_state.setdefault("draft_gender", "Not specified")
st.session_state.setdefault("draft_race", "Not specified")
st.session_state.setdefault("draft_prev_notes", "")
st.session_state.setdefault("draft_exists", False)

# Cache storage
st.session_state.setdefault("enable_pubmed", False)
st.session_state.setdefault("pubmed_results", [])
st.session_state.setdefault("word_file", None)

# Results for Tab 1 & Tab 2
st.session_state.setdefault("show_text_results", False)
st.session_state.setdefault("text_cleaned", "")
st.session_state.setdefault("text_diseases", [])
st.session_state.setdefault("text_soap_dict", {})
st.session_state.setdefault("text_post", {})

st.session_state.setdefault("show_form_results", False)
st.session_state.setdefault("form_cleaned", "")
st.session_state.setdefault("form_diseases", [])
st.session_state.setdefault("form_soap_dict", {})
st.session_state.setdefault("form_post", {})

# ---------- 3. Callback Functions (Reset & Demo) ----------
def reset_patient_data_callback():
    st.session_state.processed = False
    st.session_state.show_text_results = False
    st.session_state.show_form_results = False
    st.session_state.input_text = ""
    st.session_state.tab1_text_area = ""
    st.session_state.patient_name = ""
    st.session_state.age = "Not specified"
    st.session_state.gender = "Not specified"
    st.session_state.race = "Not specified"
    st.session_state.prev_notes = ""
    st.session_state.draft_exists = False
    st.session_state.word_file = None
    st.session_state.pubmed_results = []

    st.session_state.sb_patient_name = ""
    st.session_state.sb_age = "Not specified"
    st.session_state.sb_gender = "Not specified"
    st.session_state.sb_race = "Not specified"
    st.session_state.sb_prev_notes = ""

    st.session_state.file_version += 1
    
    # Clear Extracted Text
    st.session_state.pop("tab1_doc_preview", None)
    st.session_state.pop("tab1_stt_text", None)

    # 重置 Tab 2 表单所有输入字段
    st.session_state.f_cough_dur = ""
    st.session_state.f_cough_type = "Dry cough"
    st.session_state.f_sputum_color = "Not specified"
    st.session_state.f_sputum_vol = "Small"
    st.session_state.f_has_fever = False
    st.session_state.f_fever_temp = 38.5
    st.session_state.f_sore_throat = False
    st.session_state.f_dyspnoea_status = "Not specified"
    st.session_state.f_chest_pain_status = "Not specified"
    st.session_state.f_smoking = ""
    st.session_state.f_allergies = ""
    st.session_state.f_past_hx = ""
    st.session_state.f_current_meds = ""

def load_demo_callback():
    demo_text = """Patient Name: Tan Ah Kow, IC: 850101-13-5221, Phone: 012-3456789
Patient is a 41-year-old male presenting with severe cough for 4 days and sore throat.
Reports fever with peak temperature of 38.8°C yesterday. Cough is productive with yellow purulent sputum.
Denies chest pain, but reports slight dyspnoea when climbing stairs.
Current medications: Taking Paracetamol 500mg PRN. Smoker 10 pack-years. NKDA."""

    # 自动切回 "Text Input" 模式，解决在文档/音频切页下按 Load Demo 提示 "Please provide input text" 的问题
    st.session_state.tab1_input_mode = "Text Input"
    st.session_state.tab1_text_area = demo_text
    st.session_state.sb_patient_name = "Tan Ah Kow"
    st.session_state.sb_age = "31-50"
    st.session_state.sb_gender = "Male"
    st.session_state.sb_race = "Chinese"

    st.session_state.input_text = demo_text
    st.session_state.patient_name = "Tan Ah Kow"
    st.session_state.age = "31-50"
    st.session_state.gender = "Male"
    st.session_state.race = "Chinese"

def insert_symbol_callback(symbol):
    st.session_state.tab1_text_area = (st.session_state.tab1_text_area or "") + f" {symbol} "

# ============================================================
# app.py - Helper Functions & Display Engines
# ============================================================

# ---------- 4. PDF & OCR Extraction Helper ----------
def extract_text_from_pdf(uploaded_file):
    raw_text = ""
    method_used = "none"
    try:
        from pypdf import PdfReader
        uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)
        for page in reader.pages:
            p_text = page.extract_text()
            if p_text: raw_text += p_text + "\n"
        method_used = "pypdf"
    except Exception: pass

    if len(raw_text.strip()) < 50:
        try:
            import pdfplumber
            uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    p_text = page.extract_text()
                    if p_text: raw_text += p_text + "\n"
            method_used = "pdfplumber"
        except Exception: pass

    if len(raw_text.strip()) < 50:
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            uploaded_file.seek(0)
            images = convert_from_bytes(uploaded_file.getvalue(), dpi=150)
            ocr_text = ""
            for i, img in enumerate(images):
                ocr_text += f"--- Page {i+1} ---\n{pytesseract.image_to_string(img)}\n"
            if len(ocr_text.strip()) > len(raw_text.strip()):
                raw_text, method_used = ocr_text, "ocr"
        except Exception: pass

    msg = f"✅ Extracted {len(raw_text)} characters using {method_used}." if raw_text else "❌ Extraction failed."
    return raw_text, method_used, msg

# ---------- 5. Word Document Export (Uses soap_dict directly) ----------
def generate_word_doc(cleaned, soap_dict, missing_info, red_flags, suggestions):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading("AI Clinical Assistant - SOAP Note", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        doc.add_heading("Sanitized Patient Data", level=1)
        doc.add_paragraph(extract_demographics(cleaned))

        doc.add_heading("Medications", level=1)
        doc.add_paragraph(soap_dict.get('med_list', 'None documented'))

        doc.add_heading("SOAP Note", level=1)
        doc.add_paragraph(f"Subjective: {soap_dict.get('subjective', 'N/A')}")
        doc.add_paragraph(f"Objective: {soap_dict.get('objective', 'N/A')}")
        doc.add_paragraph(f"Assessment: {soap_dict.get('assessment', 'N/A')}")
        doc.add_paragraph(f"Plan: {soap_dict.get('plan', 'N/A')}")

        doc.add_heading("Missing Information", level=1)
        if missing_info:
            for item in missing_info: doc.add_paragraph(item, style='List Bullet')
        else: doc.add_paragraph("None")

        doc.add_heading("Red Flags & Escalation Triggers", level=1)
        if red_flags:
            for flag in red_flags: doc.add_paragraph(flag, style='List Bullet')
        else: doc.add_paragraph("None detected")

        doc.add_heading("Clinical Suggestions & Evidence References", level=1)
        for item in suggestions:
            p = doc.add_paragraph()
            p.add_run(f"{item['suggestion']} ").bold = True
            p.add_run(f"(Reference: {item['reference']})")

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return None

# ---------- 6. PubMed Search Helper ----------
@st.cache_data(ttl=3600)
def search_pubmed(query, max_results=3, tool="ai_clinical_assistant", email="clinical_app@example.com"):
    import requests
    search_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    params = {
        "db": "pubmed",
        "term": query,
        "retmax": max_results,
        "retmode": "json",
        "tool": tool,
        "email": email
    }
    try:
        resp = requests.get(search_url, params=params, timeout=10).json()
        id_list = resp.get("esearchresult", {}).get("idlist", [])
        if not id_list: return []

        sum_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"
        sum_params = {
            "db": "pubmed",
            "id": ",".join(id_list),
            "retmode": "json",
            "tool": tool,
            "email": email
        }
        sum_resp = requests.get(sum_url, params=sum_params, timeout=10).json()

        results = []
        for uid in id_list:
            article = sum_resp.get("result", {}).get(uid, {})
            results.append({
                "title": article.get("title", "No title"),
                "source": article.get("source", "Unknown"),
                "date": article.get("pubdate", "N/A"),
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{uid}/"
            })
        return results
    except Exception: return []

# ---------- 7. Shared Results Display Function (Safe HTML & Linebreaks) ----------
def extract_demographics(cleaned_text):
    """按 Key 唯一提取人口学与脱敏标识，彻底消除重复输出，按行整齐排列"""
    if not cleaned_text:
        return cleaned_text

    keys = [
        ("Patient Name", r'(?:Patient\s*Name|Patient|Name)\s*:\s*\[PATIENT_NAME\]|\[PATIENT_NAME\]'),
        ("Age Group", r'(?:Age\s*Group|Age)\s*:\s*([^\n,]+)'),
        ("Gender", r'Gender\s*:\s*([^\n,]+)'),
        ("Race", r'Race\s*:\s*([^\n,]+)'),
        ("MYKAD", r'\[MYKAD\]'),
        ("Phone", r'\[PHONE\]'),
        ("Email", r'\[EMAIL\]')
    ]

    extracted = []
    seen_keys = set()

    for label, pattern in keys:
        match = re.search(pattern, cleaned_text, re.IGNORECASE)
        if match and label not in seen_keys:
            seen_keys.add(label)
            if label == "Patient Name":
                extracted.append("Patient Name: [PATIENT_NAME]")
            elif label in ["MYKAD", "Phone", "Email"]:
                extracted.append(f"{label}: {match.group(0)}")
            else:
                extracted.append(f"{label}: {match.group(1).strip()}")

    return '\n'.join(extracted) if extracted else "No explicit demographic metadata provided."

def display_results(cleaned, diseases, soap_dict, post):
    st.markdown("---")
    st.header("📊 Clinical Output & Analysis Results")

    if post.get('red_flags'):
        st.error("### 🚨 HIGH PRIORITY RED FLAGS DETECTED")
        for flag in post['red_flags']:
            st.error(f"• **{flag}** - Requires immediate clinical evaluation!")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("🛡️ Sanitized Patient Data")
        st.caption("🔒 PII removed (Names, ICs redacted) · Age generalized")
        demo_only = extract_demographics(cleaned)
        safe_cleaned = html.escape(demo_only).replace("\n", "<br>")
        st.markdown(f'<div class="card">{safe_cleaned}</div>', unsafe_allow_html=True)

        st.subheader("🧬 Detected Clinical Entities")
        if diseases:
            st.success("Diseases identified: " + ", ".join(diseases))
        else:
            st.info("No specific disease entities isolated.")

    with col2:
        st.subheader("💊 Extracted Medications")
        safe_meds = html.escape(soap_dict.get("med_list", "- None documented")).replace("\n", "<br>")
        st.markdown(f'<div class="card">{safe_meds}</div>', unsafe_allow_html=True)

        st.subheader("📋 Generated SOAP Note")
        subj = html.escape(soap_dict.get('subjective', 'N/A')).replace("\n", "<br>")
        obj = html.escape(soap_dict.get('objective', 'N/A')).replace("\n", "<br>")
        assess = html.escape(soap_dict.get('assessment', 'N/A')).replace("\n", "<br>")
        plan_text = html.escape(soap_dict.get('plan', 'N/A')).replace("\n", "<br>")

        soap_html = f"""
        <div class="card">
            <p><strong>Subjective:</strong><br>{subj}</p>
            <p><strong>Objective:</strong><br>{obj}</p>
            <p><strong>Assessment:</strong><br>{assess}</p>
            <p><strong>Plan:</strong><br>{plan_text}</p>
        </div>
        """
        st.markdown(soap_html, unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("🔍 Clinical Decision Support & Safety Triggers")
    c1, c2, c3 = st.columns(3)

    with c1:
        st.markdown("**⚠️ Missing Clinical Information**")
        if post.get('missing_info'):
            for item in post['missing_info']: st.warning(f"• {item}")
        else: st.success("All critical elements documented.")

    with c2:
        st.markdown("**🚨 Escalation Triggers Summary**")
        if post.get('red_flags'):
            for flag in post['red_flags']: st.error(f"🚨 {flag}")
        else: st.info("No immediate red flags detected.")

    with c3:
        st.markdown("**💡 Clinical Suggestions**")
        for item in post.get('suggestions', []):
            st.info(f"**{item['suggestion']}**")

    with st.expander("📖 View All Cited Medical References & Guidelines", expanded=True):
        if st.session_state.get('enable_pubmed', False) and st.session_state.get('pubmed_results'):
            st.markdown("#### 🌐 Live PubMed Results")
            for idx, item in enumerate(st.session_state.pubmed_results):
                st.markdown(f"**{idx+1}. {item['title']}** ({item['source']}, {item['date']}) - [Link]({item['url']})")

        suggestions = post.get('suggestions', [])
        if suggestions:
            st.markdown("#### 📚 Static Evidence References")
            for idx, item in enumerate(suggestions):
                st.markdown(f"**{idx+1}. {item['suggestion']}**\n\n📚 *Source:* {item['reference']}")

    st.markdown("---")
    st.subheader("💾 Export Documentation")

    final_text = f"""=== AI Clinical Assistant - SOAP Note ===
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}

[Sanitized Patient Data]
{cleaned}

[Extracted Medications]
{soap_dict.get('med_list', 'None')}

[SOAP Note]
Subjective: {soap_dict.get('subjective', 'N/A')}
Objective: {soap_dict.get('objective', 'N/A')}
Assessment: {soap_dict.get('assessment', 'N/A')}
Plan: {soap_dict.get('plan', 'N/A')}

[Red Flags]
{', '.join(post.get('red_flags', [])) if post.get('red_flags') else 'None'}
"""
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        st.download_button("📥 Download Plain Text (.txt)", final_text, f"SOAP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt", use_container_width=True)
    with col_d2:
        if st.session_state.get('word_file'):
            st.download_button("📥 Download Word Document (.docx)", st.session_state.word_file, f"SOAP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

# ============================================================
# app.py - Draft Recovery & Sidebar Gateway
# ============================================================

# ---------- 8. Draft Recovery Banner Block ----------
if st.session_state.draft_exists and not st.session_state.processed:
    st.markdown('<div class="recovery-banner">💾 <strong>Draft found</strong> in this session.</div>', unsafe_allow_html=True)
    col_rec1, col_rec2 = st.columns(2)
    with col_rec1:
        if st.button("🔄 Restore Draft", use_container_width=True, key="btn_restore_draft"):
            st.session_state.input_text = st.session_state.draft_input
            st.session_state.patient_name = st.session_state.draft_patient_name
            st.session_state.age = st.session_state.draft_age
            st.session_state.gender = st.session_state.draft_gender
            st.session_state.race = st.session_state.draft_race
            st.session_state.prev_notes = st.session_state.draft_prev_notes

            st.session_state.sb_patient_name = st.session_state.draft_patient_name
            st.session_state.sb_age = st.session_state.draft_age
            st.session_state.sb_gender = st.session_state.draft_gender
            st.session_state.sb_race = st.session_state.draft_race
            st.session_state.sb_prev_notes = st.session_state.draft_prev_notes

            st.session_state.draft_exists = False
            st.rerun()
    with col_rec2:
        if st.button("🗑️ Discard Draft", use_container_width=True, key="btn_discard_draft"):
            st.session_state.draft_exists = False
            st.rerun()

# ---------- 9. Sidebar Input Gateway ----------
with st.sidebar:
    st.header("📥 Input Gateway")
    st.subheader("Patient Demographics")

    age_options = ["Not specified", "18-30", "31-50", "51-70", "70+"]
    curr_age_idx = age_options.index(st.session_state.age) if st.session_state.age in age_options else 0

    gender_options = ["Not specified", "Male", "Female", "Other"]
    curr_gender_idx = gender_options.index(st.session_state.gender) if st.session_state.gender in gender_options else 0

    race_options = ["Not specified", "Malay", "Chinese", "Indian", "Sarawak Native", "Sabah Native", "Others"]
    curr_race_idx = race_options.index(st.session_state.race) if st.session_state.race in race_options else 0

    st.session_state.patient_name = st.text_input("Patient Name / ID", value=st.session_state.patient_name, key="sb_patient_name")

    col_sb1, col_sb2 = st.columns(2)
    with col_sb1:
        st.session_state.age = st.selectbox("Age Group", age_options, index=curr_age_idx, key="sb_age")
    with col_sb2:
        st.session_state.gender = st.selectbox("Gender", gender_options, index=curr_gender_idx, key="sb_gender")

    st.session_state.race = st.selectbox("Race / Ethnicity", race_options, index=curr_race_idx, key="sb_race")

    st.session_state.prev_notes = st.text_area("Previous Visit Notes", value=st.session_state.prev_notes or "", height=100, key="sb_prev_notes")

    st.markdown("---")
    st.session_state.enable_pubmed = st.checkbox("🌐 Enable Live PubMed Search", value=st.session_state.enable_pubmed, key="sb_enable_pubmed")
    st.markdown("---")

    col_btn1, col_btn2 = st.columns(2)
    with col_btn1:
        st.button("📝 New Record", use_container_width=True, key="sb_new_record", on_click=reset_patient_data_callback)
    with col_btn2:
        if st.button("💾 Save Draft", use_container_width=True, key="sb_save_draft"):
            st.session_state.draft_input = st.session_state.input_text
            st.session_state.draft_patient_name = st.session_state.patient_name
            st.session_state.draft_age = st.session_state.age
            st.session_state.draft_gender = st.session_state.gender
            st.session_state.draft_race = st.session_state.race
            st.session_state.draft_prev_notes = st.session_state.prev_notes
            st.session_state.draft_exists = True
            st.success("✅ Draft saved!")

# ============================================================
# app.py - Tab 1: Unstructured Input & Symbols Bar
# ============================================================

st.title("🩺 AI Clinical Assistant")
st.caption("Your Trusted Clinical Assistant · Local‑first Medical AI")

tab1, tab2 = st.tabs(["📝 Unstructured Consultation Input", "📋 Structured Form Input"])

# ============================================================
# TAB 1: Unstructured Consultation Input
# ============================================================
with tab1:
    st.markdown("### 📝 Unstructured Consultation Notes / Dialogue")

    st.button("📋 Load Demo Consultation Record (With PII Test)", key="btn_load_demo", on_click=load_demo_callback, type="secondary")

    input_mode = st.radio("Input Method:", ["Text Input", "Upload Document (TXT/PDF)", "Upload Audio (Simulated STT)"], horizontal=True, key="tab1_input_mode")

    raw_text = ""

    if input_mode == "Text Input":
        raw_text = st.text_area("Paste Consultation Text:", height=200, key="tab1_text_area", placeholder="e.g., Patient presents with fever 38.5°C...")
        st.session_state.input_text = raw_text

        with st.container(border=True):
            st.caption("⚡ **Medical Symbols Quick-Bar** (Click to insert):")
            sym_cols = st.columns(7)
            symbols = ["°C", "SpO₂", "μg", "≥", "≤", "PRN", "NKDA"]
            for idx, sym in enumerate(symbols):
                with sym_cols[idx]:
                    st.button(sym, key=f"btn_sym_{sym}", on_click=insert_symbol_callback, args=(sym,), type="secondary")

        st.markdown("<div style='margin-bottom: 12px;'></div>", unsafe_allow_html=True)

    elif input_mode == "Upload Document (TXT/PDF)":
        uploaded_file = st.file_uploader("Upload Document", type=["txt", "pdf"], key=f"tab1_doc_upload_{st.session_state.file_version}")
        if uploaded_file:
            if uploaded_file.type == "text/plain":
                raw_text = uploaded_file.read().decode("utf-8")
            else:
                raw_text, _, msg = extract_text_from_pdf(uploaded_file)
                st.info(msg)
            st.text_area("Extracted Text Preview:", raw_text, height=180, key="tab1_doc_preview")
            st.session_state.input_text = raw_text

    else:
        audio_file = st.file_uploader("Upload Consultation Audio", type=["wav", "mp3"], key=f"tab1_audio_{st.session_state.file_version}")
        if audio_file:
            st.success("✅ Audio uploaded successfully!")
            raw_text = """Doctor: Good morning, what brings you in today?
Patient: My name is John Smith (IC: 900215-14-5533). I've had a bad cough for 3 days and fever 38.5°C.
Doctor: Are you coughing up any sputum?
Patient: Yes, yellow purulent sputum.
Doctor: Any chest pain or shortness of breath?
Patient: No chest pain, but slight dyspnoea when walking upstairs.
Doctor: Are you allergic to any medicines?
Patient: NKDA."""
            st.text_area("Transcribed Consultation Audio:", raw_text, height=180, key="tab1_stt_text")
            st.session_state.input_text = raw_text

    full_input = raw_text
    demographics_prefix = []
    if st.session_state.patient_name and not re.search(r'(?:Patient\s*Name|Name)\s*:', raw_text, re.IGNORECASE):
        demographics_prefix.append(f"Patient Name: {st.session_state.patient_name}")
    if st.session_state.age != "Not specified" and not re.search(r'Age\s*:', raw_text, re.IGNORECASE):
        demographics_prefix.append(f"Age Group: {st.session_state.age}")
    if st.session_state.gender != "Not specified" and not re.search(r'Gender\s*:', raw_text, re.IGNORECASE):
        demographics_prefix.append(f"Gender: {st.session_state.gender}")
    if st.session_state.race != "Not specified" and not re.search(r'Race\s*:', raw_text, re.IGNORECASE):
        demographics_prefix.append(f"Race: {st.session_state.race}")

    if demographics_prefix: full_input = "\n".join(demographics_prefix) + "\n\n" + full_input
    if st.session_state.prev_notes.strip():
        full_input += f"\n\n=== PREVIOUS VISIT NOTES ===\n{st.session_state.prev_notes}\n=== END PREVIOUS NOTES ===\n"

    if st.button("🚀 Run Pipeline", type="primary", key="tab1_run_btn"):
        if raw_text.strip():
            with st.spinner("🔄 Generating SOAP note..."):
                try:
                    res = run_pipeline(full_input, openmed_offline=False)
                except Exception:
                    res = run_pipeline(full_input, openmed_offline=True)

                st.session_state.text_cleaned = res["cleaned"]
                st.session_state.text_diseases = res["diseases"]
                st.session_state.text_soap_dict = res["soap_dict"]
                st.session_state.text_post = res["post"]

                st.session_state.show_text_results = True
                st.session_state.show_form_results = False

                if st.session_state.get('enable_pubmed', False) and res["diseases"]:
                    st.session_state.pubmed_results = search_pubmed(" ".join(res["diseases"][:3]))

                st.session_state.word_file = generate_word_doc(
                    res["cleaned"], res["soap_dict"],
                    res["post"].get('missing_info', []),
                    res["post"].get('red_flags', []),
                    res["post"].get('suggestions', [])
                )
                st.success("✅ Complete!")
        else:
            st.warning("⚠️ Please provide input text.")

    if st.session_state.get('show_text_results', False):
        display_results(
            st.session_state.text_cleaned,
            st.session_state.text_diseases,
            st.session_state.text_soap_dict,
            st.session_state.text_post
        )

# ============================================================
# app.py - Tab 2: Structured Form Container (st.form)
# ============================================================

# ============================================================
# TAB 2: Structured Form Input
# ============================================================
with tab2:
    st.markdown("### 📋 Respiratory Consultation Guided Form")

    with st.container(border=True):
        col_form1, col_form2 = st.columns(2)

        with col_form1:
            with st.container(border=True):
                st.markdown("**1. Chief Complaint & Cough**")
                cough_duration = st.text_input("Cough Duration (e.g., 3 days)", key="f_cough_dur")
                cough_type = st.radio("Cough Type:", ["Dry cough", "Productive cough (with sputum)"], horizontal=True, key="f_cough_type")

            with st.container(border=True):
                st.markdown("**2. Sputum Characteristics**")
                sputum_color = st.selectbox("Sputum Color:", ["Not specified", "White/Clear", "Yellow purulent", "Rust-colored", "Blood-tinged"], key="f_sputum_color")
                sputum_volume = st.radio("Sputum Volume:", ["Small", "Moderate", "Large"], horizontal=True, key="f_sputum_vol")

            with st.container(border=True):
                st.markdown("**3. Fever & Vitals**")
                has_fever = st.checkbox("Fever present", key="f_has_fever")
                fever_temp = st.number_input("Peak Temperature (°C)", min_value=35.0, max_value=42.0, value=38.5, step=0.1, key="f_fever_temp")

        with col_form2:
            with st.container(border=True):
                st.markdown("**4. Associated Symptoms & Negatives**")
                col_sx1, col_sx2 = st.columns(2)
                with col_sx1:
                    has_sore_throat = st.checkbox("Sore throat", key="f_sore_throat")
                    # Dyspnoea 三选一
                    dyspnoea_status = st.radio(
                        "Dyspnoea / Shortness of breath:",
                        ["Not specified", "Present", "Denied"],
                        horizontal=True,
                        key="f_dyspnoea_status"
                    )
                with col_sx2:
                    # Chest Pain 三选一
                    chest_pain_status = st.radio(
                        "Chest pain:",
                        ["Not specified", "Present", "Denied"],
                        horizontal=True,
                        key="f_chest_pain_status"
                    )

            with st.container(border=True):
                st.markdown("**5. Personal History**")
                smoking = st.text_input("Smoking History", key="f_smoking")
                allergies = st.text_input("Allergies", key="f_allergies")

        with st.container(border=True):
            st.markdown("**6. Past Medical History & Current Medications**")
            col_add1, col_add2 = st.columns(2)
            with col_add1:
                past_history = st.text_area("Past Medical History", key="f_past_hx", height=70)
            with col_add2:
                medications = st.text_area("Current Medications", key="f_current_meds", height=70)

        form_submitted = st.button("🩺 Generate SOAP Note from Form", type="primary", key="btn_form_submit")

    # --- 表单提交后的处理逻辑 ---
    if form_submitted:
        narrative_parts = []
        is_dry = (cough_type == "Dry cough")
        cough_desc = "dry cough" if is_dry else "productive cough"

        narrative_parts.append(f"a {cough_desc} for {cough_duration}" if cough_duration else f"a {cough_desc}")

        if not is_dry:
            if sputum_color != "Not specified":
                narrative_parts.append(f"{sputum_color.lower()} sputum")
        if has_fever:
            narrative_parts.append(f"fever {fever_temp:.1f}°C")

        # 构建症状列表（三选一）
        sx_list = []
        if has_sore_throat:
            sx_list.append("sore throat")

        if dyspnoea_status == "Present":
            sx_list.append("dyspnoea")
        elif dyspnoea_status == "Denied":
            sx_list.append("denies dyspnoea")

        if chest_pain_status == "Present":
            sx_list.append("chest pain")
        elif chest_pain_status == "Denied":
            sx_list.append("denies chest pain")

        if sx_list:
            narrative_parts.append(f"associated with {', '.join(sx_list)}")

        narrative = "Patient presents with " + ", ".join(narrative_parts) + "."
        if past_history:
            narrative += f" Past history: {past_history}."
        if medications:
            narrative += f" Current medications: {medications}."
        if allergies:
            narrative += f" Allergies: {allergies}."
        if smoking:
            narrative += f" Smoker: {smoking}."

        # 添加侧边栏人口学信息
        demographics_prefix = []
        if st.session_state.patient_name and not re.search(r'(?:Patient\s*Name|Name)\s*:', narrative, re.IGNORECASE):
            demographics_prefix.append(f"Patient Name: {st.session_state.patient_name}")
        if st.session_state.age != "Not specified" and not re.search(r'Age\s*:', narrative, re.IGNORECASE):
            demographics_prefix.append(f"Age Group: {st.session_state.age}")
        if st.session_state.gender != "Not specified" and not re.search(r'Gender\s*:', narrative, re.IGNORECASE):
            demographics_prefix.append(f"Gender: {st.session_state.gender}")
        if st.session_state.race != "Not specified" and not re.search(r'Race\s*:', narrative, re.IGNORECASE):
            demographics_prefix.append(f"Race: {st.session_state.race}")
        if demographics_prefix:
            narrative = "\n".join(demographics_prefix) + "\n\n" + narrative

        if st.session_state.prev_notes.strip():
            narrative += f"\n\n=== PREVIOUS VISIT NOTES ===\n{st.session_state.prev_notes}\n=== END PREVIOUS NOTES ===\n"

        with st.spinner("🔄 Generating SOAP note from form..."):
            try:
                res = run_pipeline(narrative, openmed_offline=False)
            except Exception:
                res = run_pipeline(narrative, openmed_offline=True)

            st.session_state.form_cleaned = res["cleaned"]
            st.session_state.form_diseases = res["diseases"]
            st.session_state.form_soap_dict = res["soap_dict"]
            st.session_state.form_post = res["post"]

            st.session_state.show_form_results = True
            st.session_state.show_text_results = False

            st.session_state.word_file = generate_word_doc(
                res["cleaned"], res["soap_dict"],
                res["post"].get('missing_info', []),
                res["post"].get('red_flags', []),
                res["post"].get('suggestions', [])
            )
            st.success("✅ SOAP note generated successfully!")

    if st.session_state.get('show_form_results', False):
        display_results(
            st.session_state.form_cleaned,
            st.session_state.form_diseases,
            st.session_state.form_soap_dict,
            st.session_state.form_post
        )
